"""
文档标题格式化器 - 负责文档标题和附件处理
"""

from docx.document import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

from .base_formatter import BaseFormatter


class DocumentTitleFormatter(BaseFormatter):
    """文档标题格式化器 - 负责文档标题和附件处理"""

    def add_document_title(self, doc: Document, title: str):
        """在文档开头添加标题"""
        # Create paragraphs through the public API, then move their XML nodes to
        # the start so documents containing only tables are supported as well.
        title_paragraph = doc.add_paragraph()
        blank_paragraph = doc.add_paragraph()
        body = doc._body._element
        body.insert(0, title_paragraph._element)
        body.insert(1, blank_paragraph._element)
        title_paragraph.alignment = self.config.ALIGNMENTS['center']

        run = title_paragraph.add_run(title)
        run.font.size = self.config.FONT_SIZES['title']
        run.bold = False

        self._set_run_fonts(run, self.config.FONTS['xiaobiaosong'])

        # 设置段落格式，确保与网格对齐
        paragraph_format = title_paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = self.config.TITLE_LINE_PITCH
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # A size-two title has its own line definition and must not be forced
        # onto the size-three body grid.
        self._disable_snap_to_grid(title_paragraph)

        blank_format = blank_paragraph.paragraph_format
        blank_format.space_before = Pt(0)
        blank_format.space_after = Pt(0)
        self._enable_snap_to_grid(blank_paragraph)
