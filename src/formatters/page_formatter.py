"""
页面格式化器 - 负责页面设置、页眉页脚、页码等
"""

from docx.document import Document
from docx.oxml.ns import qn as qn_func
from docx.oxml.shared import OxmlElement, qn
from docx.section import Section
from docx.shared import Emu, Mm, Pt

from .base_formatter import BaseFormatter


class PageFormatter(BaseFormatter):
    """页面格式化器 - 负责页面设置、页眉页脚、页码等"""

    def setup_page_format(self, doc: Document) -> None:
        """设置页面格式"""
        for section in doc.sections:
            self._setup_section_format(section)

        # 清零文档级默认段前段后间距，避免未被格式化器覆盖的空段落继承 after=200
        self._zero_default_paragraph_spacing(doc)

    def _setup_section_format(self, section: Section) -> None:
        """Apply the configured page geometry to one document section."""
        section.orientation = self.config.PAGE_ORIENTATION
        section.page_width = self.config.PAGE_SIZE['width']
        section.page_height = self.config.PAGE_SIZE['height']

        section.top_margin = self.config.PAGE_MARGINS['top']
        section.bottom_margin = self.config.PAGE_MARGINS['bottom']
        section.left_margin = self.config.PAGE_MARGINS['left']
        section.right_margin = self.config.PAGE_MARGINS['right']

        section.header_distance = Mm(25)
        section.footer_distance = Mm(25)

        self._setup_document_grid(section)

    def _zero_default_paragraph_spacing(self, doc: Document):
        """将 styles.xml 中的 docDefaults 段前段后清零。

        Pandoc 默认会在 docDefaults 中写入 <w:spacing w:after="200"/>，
        导致未被任一格式化器显式设置的段落（如标题后插入的空行）保留 10pt 间距。
        """
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn_func('w:docDefaults'))
        if doc_defaults is None:
            return
        p_pr_default = doc_defaults.find(qn_func('w:pPrDefault'))
        if p_pr_default is None:
            return
        p_pr = p_pr_default.find(qn_func('w:pPr'))
        if p_pr is None:
            return
        spacing = p_pr.find(qn_func('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            p_pr.append(spacing)
        spacing.set(qn_func('w:before'), '0')
        spacing.set(qn_func('w:after'), '0')
        # 防御性清除自动间距（若启用会忽略显式值）
        spacing.set(qn_func('w:beforeAutospacing'), '0')
        spacing.set(qn_func('w:afterAutospacing'), '0')

    def add_page_numbers(self, doc: Document):
        """Add mirrored odd/even page numbers required by the official format."""
        doc.settings.odd_and_even_pages_header_footer = True

        for section in doc.sections:
            self._configure_page_number_footer(
                section.footer,
                self.config.ALIGNMENTS['right'],
                right_indent=Pt(16),
            )
            self._configure_page_number_footer(
                section.even_page_footer,
                self.config.ALIGNMENTS['left'],
                left_indent=Pt(16),
            )

    def _configure_page_number_footer(
        self,
        footer,
        alignment,
        left_indent=None,
        right_indent=None,
    ):
        """Configure one footer with a PAGE field and one-character side margins."""
        paragraph = footer.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = alignment
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = left_indent
        paragraph_format.right_indent = right_indent
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        leading_run = paragraph.add_run('— ')
        self._format_page_number_run(leading_run)

        field_run = paragraph.add_run()
        field_begin = OxmlElement('w:fldChar')
        field_begin.set(qn('w:fldCharType'), 'begin')
        instruction = OxmlElement('w:instrText')
        instruction.set(qn('xml:space'), 'preserve')
        instruction.text = ' PAGE '
        field_separator = OxmlElement('w:fldChar')
        field_separator.set(qn('w:fldCharType'), 'separate')
        cached_value = OxmlElement('w:t')
        cached_value.text = '1'
        field_end = OxmlElement('w:fldChar')
        field_end.set(qn('w:fldCharType'), 'end')
        for element in (field_begin, instruction, field_separator, cached_value, field_end):
            field_run._r.append(element)
        self._format_page_number_run(field_run)

        trailing_run = paragraph.add_run(' —')
        self._format_page_number_run(trailing_run)

    def _format_page_number_run(self, run):
        """Apply the configured Chinese size-four Song typeface to a page-number run."""
        run.font.size = self.config.FONT_SIZES['page_num']
        self._set_run_fonts(run, self.config.FONTS['songti'])

    def _setup_document_grid(self, section):
        """设置文档网格以强制每页22行、每行28字"""
        sectPr = section._sectPr

        # 检查是否已有docGrid元素
        docGrid = sectPr.find(qn_func('w:docGrid'))
        if docGrid is None:
            # 创建新的docGrid元素
            docGrid = OxmlElement('w:docGrid')
            sectPr.append(docGrid)

        # 设置文档网格类型为"行和字符网格"
        docGrid.set(qn_func('w:type'), 'linesAndChars')

        # The electronic-document standard defines a body line independently
        # from the software page-margin box. Do not derive it from page height.
        line_pitch = self.config.BODY_LINE_PITCH.twips
        content_width_twips = Emu(self.config.get_content_width_emu()).twips
        char_space = round(content_width_twips / self.config.CHARS_PER_LINE)

        # 设置行距（单位：twips，1/20点）
        docGrid.set(qn_func('w:linePitch'), str(line_pitch))

        # 设置字符间距（单位：twips）
        docGrid.set(qn_func('w:charSpace'), str(char_space))
