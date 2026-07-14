"""
图片格式化器 - 负责图片处理和格式化
"""

from copy import deepcopy

from docx.document import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.shared import Pt

from ..config import DocumentConfig
from ..utils.constants import Patterns
from ..utils.exceptions import ImageProcessingError, XMLProcessingError
from ..utils.xpath_cache import OptimizedXMLProcessor
from .base_formatter import BaseFormatter


class ImageFormatter(BaseFormatter):
    """图片格式化器 - 负责图片处理和格式化"""

    def __init__(self, config: DocumentConfig | None = None):
        super().__init__(config)
        self.xml_processor = OptimizedXMLProcessor()

    def format_images(self, doc: Document):
        """格式化文档中的图片 - 使用优化的单次遍历，并移除所有图片相关文件名"""
        try:
            # 1. 使用优化的方法一次性查找所有包含图片的段落
            drawings_map = self.xml_processor.find_drawings_in_paragraphs(doc.paragraphs)

            # 处理每个包含图片的段落
            for paragraph_index, drawings in drawings_map.items():
                for drawing in drawings:
                    self._format_single_image(drawing)
                self._collapse_anchored_image_paragraph(doc.paragraphs[paragraph_index])

            # 2. 格式化图片标题段落（处理"Image Caption"样式和"图 X："格式）
            self._format_all_image_captions(doc)

        except (AttributeError, KeyError) as error:
            raise ImageProcessingError(f'图片元素访问错误: {error}') from error
        except Exception as error:
            raise ImageProcessingError(f'图片格式化时出现错误: {error}') from error

    def _format_single_image(self, drawing_element):
        """格式化单个图片"""
        self._set_image_full_width(drawing_element)

        if self.config.PANDOC_CONFIG.get('image_wrap_text', False):
            self._set_image_wrap(drawing_element)

    def _collapse_anchored_image_paragraph(self, paragraph):
        """Collapse the anchor paragraph so it adds no visible line before a caption."""
        if paragraph.text.strip() or not paragraph._element.xpath('.//wp:anchor'):
            return

        paragraph.alignment = self.config.ALIGNMENTS['center']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(1)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def _format_all_image_captions(self, doc: Document):
        """格式化所有图片标题段落（包括Pandoc生成的"Image Caption"样式）"""
        for paragraph in doc.paragraphs:
            if self._has_math_formula(paragraph):
                continue

            style = paragraph.style
            if style is not None and style.name == 'Image Caption':
                self._format_image_caption(paragraph)
                continue

            text = paragraph.text.strip()
            if text and Patterns.CAPTION_PATTERN.match(text):
                self._format_image_caption(paragraph)

    def _set_image_full_width(self, drawing_element):
        """设置图片宽度为全宽（适用于inline元素）"""
        try:
            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)

            content_width_emu = self.config.get_content_width_emu()

            # 获取并更新extent元素
            for extent in elements.get('extent', []):
                # 获取原始尺寸
                cx_original = extent.get('cx', '3000000')
                cy_original = extent.get('cy', '2000000')

                # 计算纵横比并根据新宽度计算高度
                try:
                    cx_orig_int = int(cx_original)
                    cy_orig_int = int(cy_original)
                    aspect_ratio = cy_orig_int / cx_orig_int

                    # 设置图片宽度为页面宽度
                    extent.set('cx', str(content_width_emu))
                    # 按比例计算新高度
                    extent.set('cy', str(int(content_width_emu * aspect_ratio)))
                except (ValueError, ZeroDivisionError):
                    # 如果转换失败，只设置宽度
                    extent.set('cx', str(content_width_emu))

            # 同时更新pic:spPr中的extent（如果存在）
            pic_extent_xpath = './/pic:spPr/a:xfrm/a:ext'
            pic_extents = drawing_element.xpath(pic_extent_xpath)
            for pic_extent in pic_extents:
                cx_original = pic_extent.get('cx', '3000000')
                cy_original = pic_extent.get('cy', '2000000')

                try:
                    cx_orig_int = int(cx_original)
                    cy_orig_int = int(cy_original)
                    aspect_ratio = cy_orig_int / cx_orig_int

                    pic_extent.set('cx', str(content_width_emu))
                    pic_extent.set('cy', str(int(content_width_emu * aspect_ratio)))
                except (ValueError, ZeroDivisionError):
                    pic_extent.set('cx', str(content_width_emu))

        except Exception as error:
            raise ImageProcessingError(f'Failed to resize image: {error}') from error

    def _set_image_wrap(self, drawing_element):
        """设置图片文字环绕为top and bottom，并设置图片宽度为全宽"""
        try:
            if not self.config.PANDOC_CONFIG.get('image_wrap_text', False):
                return

            # 使用优化的批量查询获取所有图片相关元素
            elements = self.xml_processor.process_image_properties(drawing_element)

            for inline in elements.get('inline', []):
                # 获取父元素
                parent = inline.getparent()

                # 从批量查询结果中获取元素
                extent = elements.get('extent', [])
                cx_original = extent[0].get('cx') if extent else '3000000'
                cy_original = extent[0].get('cy') if extent else '2000000'

                content_width_emu = self.config.get_content_width_emu()

                # 计算纵横比并根据新宽度计算高度
                try:
                    cx_orig_int = int(cx_original)
                    cy_orig_int = int(cy_original)
                    aspect_ratio = cy_orig_int / cx_orig_int

                    # 设置图片宽度为页面宽度
                    cx = str(content_width_emu)
                    # 按比例计算新高度
                    cy = str(int(content_width_emu * aspect_ratio))
                except (ValueError, ZeroDivisionError):
                    # 如果转换失败，使用默认值
                    cx = str(content_width_emu)
                    cy = cy_original

                # 获取graphic元素
                graphic = elements.get('graphic', [])

                if graphic:
                    doc_properties = elements.get('docPr', [])
                    source_doc_properties = doc_properties[0] if doc_properties else None
                    anchor = self._create_anchor_element(cx, cy, source_doc_properties)

                    # 复制graphic元素到新的anchor中
                    anchor.append(graphic[0])

                    # 替换inline元素
                    parent.replace(inline, anchor)

        except Exception as error:
            raise ImageProcessingError(f'Failed to configure image wrapping: {error}') from error

    def _create_anchor_element(self, cx: str, cy: str, source_doc_properties=None):
        """
        安全创建anchor元素，避免XML注入漏洞

        Args:
            cx: 图片宽度
            cy: 图片高度
            source_doc_properties: Existing wp:docPr metadata to preserve

        Returns:
            创建的anchor XML元素
        """
        try:
            # 输入验证和清理
            cx = str(cx).strip() if cx else '3000000'
            cy = str(cy).strip() if cy else '2000000'

            # 验证数值参数
            try:
                int(cx)
                int(cy)
            except ValueError:
                # 如果参数无效，使用默认值
                cx, cy = '3000000', '2000000'

            # 使用原生XML API安全构建anchor元素
            anchor = OxmlElement('wp:anchor')

            # 设置anchor属性 (修复: 使用unqualified属性名，符合WordprocessingML规范)
            anchor.set('distT', '0')
            anchor.set('distB', '0')
            anchor.set('distL', '114300')
            anchor.set('distR', '114300')
            anchor.set('simplePos', '0')
            anchor.set('relativeHeight', '251658240')
            anchor.set('behindDoc', '0')
            anchor.set('locked', '0')
            anchor.set('layoutInCell', '1')
            anchor.set('allowOverlap', '1')

            # 创建simplePos子元素
            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0')
            simplePos.set('y', '0')
            anchor.append(simplePos)

            # 创建positionH子元素
            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'column')
            align = OxmlElement('wp:align')
            align.text = 'center'
            positionH.append(align)
            anchor.append(positionH)

            # 创建positionV子元素
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'paragraph')
            posOffset = OxmlElement('wp:posOffset')
            posOffset.text = '0'
            positionV.append(posOffset)
            anchor.append(positionV)

            # 创建extent子元素
            extent = OxmlElement('wp:extent')
            extent.set('cx', cx)
            extent.set('cy', cy)
            anchor.append(extent)

            # 创建effectExtent子元素
            effectExtent = OxmlElement('wp:effectExtent')
            effectExtent.set('l', '0')
            effectExtent.set('t', '0')
            effectExtent.set('r', '0')
            effectExtent.set('b', '0')
            anchor.append(effectExtent)

            # 创建wrapTopAndBottom子元素
            wrapTopAndBottom = OxmlElement('wp:wrapTopAndBottom')
            anchor.append(wrapTopAndBottom)

            # Preserve accessibility metadata such as descr/title while moving
            # the drawing from wp:inline to wp:anchor.
            if source_doc_properties is not None:
                docPr = deepcopy(source_doc_properties)
            else:
                docPr = OxmlElement('wp:docPr')
                docPr.set('id', '1')
                docPr.set('name', 'Picture')
            anchor.append(docPr)

            # 创建cNvGraphicFramePr子元素
            cNvGraphicFramePr = OxmlElement('wp:cNvGraphicFramePr')
            graphicFrameLocks = OxmlElement('a:graphicFrameLocks')
            graphicFrameLocks.set('noChangeAspect', '1')
            cNvGraphicFramePr.append(graphicFrameLocks)
            anchor.append(cNvGraphicFramePr)

            return anchor

        except Exception as error:
            raise XMLProcessingError(f'Failed to create image anchor: {error}') from error

    def _format_image_caption(self, paragraph):
        """格式化图片标题为caption格式"""
        for run in paragraph.runs:
            run.font.size = self.config.FONT_SIZES['table']
            run.bold = False
            self._set_run_fonts(run, self.config.FONTS['fangsong'])

        paragraph.alignment = self.config.ALIGNMENTS['center']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.first_line_indent = Pt(0)
