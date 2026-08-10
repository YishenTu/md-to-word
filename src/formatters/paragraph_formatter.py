"""
段落格式化器 - 负责标题和正文段落的格式化
"""

import re
from typing import Any

from docx.document import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from ..utils.constants import ControlTokens, Patterns
from .base_formatter import BaseFormatter


class ParagraphFormatter(BaseFormatter):
    """段落格式化器 - 负责标题和正文段落的格式化"""

    def format_document_content(self, doc: Document, metadata: dict[str, Any]):
        """格式化文档内容，应用公文格式要求"""
        self._ensure_attachment_spacing(doc)

        for paragraph in doc.paragraphs:
            # 跳过空段落
            if not paragraph.text.strip():
                continue

            # 判断段落类型并应用相应格式
            if self._is_heading(paragraph):
                self._format_heading(paragraph)
            else:
                self._format_body_paragraph(paragraph)

    def _ensure_attachment_spacing(self, doc: Document) -> None:
        """Ensure one grid-aligned blank line before every attachment declaration."""
        for paragraph in list(doc.paragraphs):
            if not self._is_attachment_first_item(paragraph):
                continue

            previous_element = paragraph._element.getprevious()
            if previous_element is not None and previous_element.tag == qn('w:p'):
                previous_paragraph = Paragraph(previous_element, paragraph._parent)
                if self._is_empty_paragraph(previous_paragraph):
                    self._format_blank_grid_line(previous_paragraph)
                    continue

            blank_paragraph = paragraph.insert_paragraph_before()
            self._format_blank_grid_line(blank_paragraph)

    @staticmethod
    def _is_attachment_first_item(paragraph: Paragraph) -> bool:
        """Return whether a paragraph starts an attachment declaration."""
        text = paragraph.text.lstrip()
        return text.startswith((ControlTokens.ATTACHMENT_FIRST_ITEM, '附件：', '附件:'))

    @staticmethod
    def _attachment_item_kind(paragraph: Paragraph) -> str | None:
        """Return the marked attachment item kind, including legacy compact declarations."""
        text = paragraph.text.lstrip()
        if text.startswith(ControlTokens.ATTACHMENT_FIRST_ITEM):
            return 'first'
        if text.startswith(ControlTokens.ATTACHMENT_ITEM):
            return 'subsequent'
        if text.startswith(('附件：', '附件:')):
            return 'first'
        return None

    @staticmethod
    def _is_empty_paragraph(paragraph: Paragraph) -> bool:
        """Return whether a paragraph has no visible text or embedded content."""
        if paragraph.text.strip():
            return False
        return not paragraph._element.xpath('.//w:drawing | .//w:br | .//m:oMath')

    def _format_blank_grid_line(self, paragraph) -> None:
        """Format an empty paragraph as exactly one document-grid line."""
        paragraph.alignment = self.config.ALIGNMENTS['left']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        self._enable_snap_to_grid(paragraph)

    def _is_heading(self, paragraph) -> bool:
        """判断段落是否为标题"""
        # 检查是否使用了标题样式
        if paragraph.style.name.startswith('Heading'):
            return True

        # 检查段落内容是否以中文数字或序号开头（如"一、"、"（一）"等）
        text = paragraph.text.strip()

        for pattern in Patterns.HEADING_PATTERNS:
            if pattern.match(text):
                return True

        return False

    def _format_heading(self, paragraph):
        """格式化标题段落"""
        text = paragraph.text.strip()
        level = self._get_heading_level(paragraph, text)

        # 如果不是一级或二级标题，按正文处理
        if level == 0:
            self._format_body_paragraph(paragraph)
            return

        # 检查是否包含数学公式（通过检查XML元素）
        if self._has_math_formula(paragraph):
            # 如果包含数学公式，只修改字体格式，不清除内容
            self._format_paragraph_with_math(paragraph, level, is_heading=True)
            return

        # 对于不包含数学公式的标题，使用原有的方法
        # 清除现有格式
        for run in paragraph.runs:
            run.clear()

        # 重新添加文本
        run = paragraph.add_run(text)

        # 根据级别应用格式
        if level == 1:
            # 一级标题：黑体，三号，不加粗
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_fonts(run, self.config.FONTS['heiti'])
        else:
            # 二级标题：楷体，三号，不加粗
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_fonts(run, self.config.FONTS['kaiti'])

        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        # GB/T 9704-2012要求：不使用段前段后间距，所有内容锁定在文档网格内
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

        # 启用文档网格对齐
        self._enable_snap_to_grid(paragraph)

    def _get_heading_level(self, paragraph, text: str) -> int:
        """获取标题级别（只处理一级和二级标题）"""
        # 检查Word内置标题样式
        if paragraph.style.name == 'Heading 1':
            return 1  # # → 黑体 (文档主标题，通常跳过)
        elif paragraph.style.name == 'Heading 2':
            return 1  # ## → 黑体
        elif paragraph.style.name == 'Heading 3':
            return 2  # ### → 楷体
        elif paragraph.style.name.startswith('Heading'):
            return 0  # #### 及以下 → 正文样式（仿宋）

        # 根据文本内容判断级别（处理中文标题格式）
        if Patterns.HEADING_PATTERNS[0].match(text):
            return 1  # 一、二、三、 → 黑体
        elif Patterns.HEADING_PATTERNS[1].match(text):
            return 2  # （一）（二）（三） → 楷体

        # 其他情况返回0，表示不是标题
        return 0

    def _format_body_paragraph(self, paragraph):
        """格式化正文段落"""
        attachment_item_kind = self._attachment_item_kind(paragraph)
        if attachment_item_kind is not None:
            self._prepare_attachment_item(paragraph, attachment_item_kind)

        # 检查是否包含数学公式
        if self._has_math_formula(paragraph):
            # 如果包含数学公式，使用特殊的格式化方法
            self._format_paragraph_with_math(paragraph, level=0, is_heading=False)
            return

        # 为所有运行应用仿宋格式
        for run in paragraph.runs:
            run.font.size = self.config.FONT_SIZES['body']
            run.bold = False
            run.italic = False
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_fonts(run, self.config.FONTS['fangsong'])

        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        if attachment_item_kind is not None:
            paragraph.alignment = self.config.ALIGNMENTS['left']
            self._set_attachment_indent(paragraph, attachment_item_kind)
        else:
            paragraph.alignment = self.config.ALIGNMENTS['justify']
            paragraph_format.left_indent = Pt(0)
            paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

        # 启用文档网格对齐
        self._enable_snap_to_grid(paragraph)

    def _prepare_attachment_item(self, paragraph: Paragraph, item_kind: str) -> None:
        """Consume the layout token and retain one marker-to-name separator."""
        text = paragraph.text.lstrip()
        if item_kind == 'first':
            text = text.removeprefix(ControlTokens.ATTACHMENT_FIRST_ITEM)
            match = re.match(r'^(附件[:：]\s*\d+\.)\s+(.+)$', text)
        else:
            text = text.removeprefix(ControlTokens.ATTACHMENT_ITEM)
            match = re.match(r'^(\d+\.)\s+(.+)$', text)

        if match is None:
            return

        marker, item_name = match.groups()
        paragraph.clear()
        paragraph.add_run(f'{marker} ')
        paragraph.add_run(item_name)

    def _set_attachment_indent(self, paragraph: Paragraph, item_kind: str) -> None:
        """Align every wrapped line with the first character of the attachment name."""
        settings = self.config.ATTACHMENT_LIST
        character_width_pt = self.config.FONT_SIZES['body'].pt
        text_position_pt = settings['text_position_chars'] * character_width_pt
        marker_position_chars = (
            settings['first_marker_position_chars']
            if item_kind == 'first'
            else settings['subsequent_marker_position_chars']
        )
        marker_position_pt = marker_position_chars * character_width_pt

        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(text_position_pt)
        paragraph_format.first_line_indent = Pt(marker_position_pt - text_position_pt)

    def _format_paragraph_with_math(self, paragraph, level: int, is_heading: bool):
        """格式化包含数学公式的段落，保留数学公式内容"""
        # 设置段落格式
        paragraph.alignment = self.config.ALIGNMENTS['justify']
        paragraph_format = paragraph.paragraph_format

        # 设置段落格式
        paragraph_format.first_line_indent = self.config.FIRST_LINE_INDENT
        paragraph_format.space_after = Pt(0)
        paragraph_format.space_before = Pt(0)

        # 启用文档网格对齐
        self._enable_snap_to_grid(paragraph)

        # 只格式化文本run，跳过数学公式
        for run in paragraph.runs:
            if run._element.tag.endswith('r'):  # 普通文本run
                try:
                    # 检查run的XML内容，只处理不包含数学内容的run
                    if 'oMath' not in run._element.xml:
                        if is_heading and level > 0:
                            # 根据标题级别设置字体（只处理一级和二级标题）
                            if level == 1:
                                self._set_run_fonts(run, self.config.FONTS['heiti'])
                            else:  # level == 2
                                self._set_run_fonts(run, self.config.FONTS['kaiti'])
                            run.font.size = self.config.FONT_SIZES['body']
                            run.bold = False
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            # 正文段落格式
                            run.font.size = self.config.FONT_SIZES['body']
                            run.bold = False
                            run.italic = False
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            self._set_run_fonts(run, self.config.FONTS['fangsong'])
                except Exception as e:
                    # 记录错误但不中断处理
                    import logging

                    logging.warning(f'处理文本段落字体时出错: {e}')
                    continue
