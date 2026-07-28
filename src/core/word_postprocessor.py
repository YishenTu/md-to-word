import os
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from ..config import DocumentConfig
from ..formatters import (
    DocumentTitleFormatter,
    ImageFormatter,
    ListFormatter,
    PageFormatter,
    ParagraphFormatter,
    SignatureFormatter,
    TableFormatter,
)
from ..utils.constants import ControlTokens, Patterns
from ..utils.exceptions import FileProcessingError, ImageProcessingError


class WordPostprocessor:
    """
    重构后的Word文档后处理器
    使用组合模式，将不同的格式化功能委托给专门的格式化器类
    解决了原有的"God Object"反模式问题

    特别处理：包含数学公式的图片caption
    - 检测包含LaTeX数学公式的caption（如 $\\theta$, $\\omega$）
    - 使用特殊的分离处理逻辑，避免破坏MathML内容
    - 通过元素属性标记已处理的caption，防止重复处理
    """

    def __init__(self):
        self.config = DocumentConfig()
        self._doc: DocxDocument | None = None

        # 初始化专门的格式化器
        self.page_formatter = PageFormatter(self.config)
        self.paragraph_formatter = ParagraphFormatter(self.config)
        self.title_formatter = DocumentTitleFormatter(self.config)
        self.table_formatter = TableFormatter(self.config)
        self.list_formatter = ListFormatter(self.config)
        self.image_formatter = ImageFormatter(self.config)
        self.signature_formatter = SignatureFormatter(self.config)

    @property
    def doc(self) -> DocxDocument:
        """Return the active document after the formatting pipeline has started."""
        if self._doc is None:
            raise FileProcessingError('Word postprocessing has no active document')
        return self._doc

    @doc.setter
    def doc(self, value: DocxDocument) -> None:
        self._doc = value

    def apply_formatting(
        self,
        docx_path: str,
        metadata: dict[str, Any],
        source_dir: str | Path | None = None,
    ) -> str:
        """
        对pandoc生成的Word文档应用公文格式

        Args:
            docx_path: pandoc生成的Word文档路径
            metadata: Document metadata containing the title
            source_dir: Markdown source directory used to resolve relative assets

        Returns:
            处理后的Word文档路径
        """
        # 加载pandoc生成的文档
        self.doc = Document(docx_path)

        # 保存文档路径的目录，用于查找图片
        self.doc_dir = os.path.dirname(os.path.abspath(docx_path))
        self.source_dir = os.path.abspath(str(source_dir)) if source_dir else self.doc_dir

        # 使用专门的格式化器处理不同方面的格式化
        self.page_formatter.setup_page_format(self.doc)

        self.paragraph_formatter.format_document_content(self.doc, metadata)

        # 添加文档标题（如果有）
        if metadata.get('title'):
            self.title_formatter.add_document_title(self.doc, metadata['title'])

        # 应用各种格式化
        self.page_formatter.add_page_numbers(self.doc)
        self.list_formatter.format_lists(self.doc)
        self.table_formatter.format_tables(self.doc)

        # 处理分页符标记
        self._process_page_breaks()

        # 新的图片处理方式：直接查找并替换图片语法
        self.process_and_insert_images()
        self.image_formatter.format_images(self.doc)

        self._process_signature(metadata)

        # 保存格式化后的文档
        self.doc.save(docx_path)
        return docx_path

    def _has_math_formula(self, paragraph) -> bool:
        """检测段落是否包含数学公式"""
        return self.image_formatter._has_math_formula(paragraph)

    def _process_page_breaks(self):
        """处理 [PAGEBREAK] 标记，转换为实际分页符"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip() == ControlTokens.PAGE_BREAK:
                # 清空段落内容
                paragraph.clear()

                # 添加分页符
                run = paragraph.add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._element.append(br)

    def _process_signature(self, metadata: dict[str, Any]) -> None:
        """Replace one signature position anchor with structured signature metadata."""
        anchors = [paragraph for paragraph in self.doc.paragraphs if paragraph.text.strip() == ControlTokens.SIGNATURE]
        signatory = metadata.get('signatory')
        document_date = metadata.get('document_date')

        if signatory is None and document_date is None:
            if anchors:
                raise FileProcessingError('Signature anchor is present without signature metadata')
            return

        if not isinstance(signatory, str) or not isinstance(document_date, str):
            raise FileProcessingError('Signature metadata must include string signatory and document_date values')
        if len(anchors) != 1:
            raise FileProcessingError(f'Expected exactly one signature anchor, found {len(anchors)}')

        self.signature_formatter.replace_signature_anchor(
            anchors[0],
            signatory,
            document_date,
        )

    def process_and_insert_images(self):
        """处理文档中的图片语法并插入实际图片"""
        # 初始化已处理的数学caption文本集合
        self._processed_math_caption_texts: set[str] = set()

        # 使用预编译的模式
        obsidian_image_pattern = Patterns.OBSIDIAN_IMAGE_PATTERN
        caption_pattern = Patterns.CAPTION_PREFIX_PATTERN

        # 收集需要处理的图片信息
        images_to_process: list[dict[str, Any]] = []

        # 第一遍：识别所有图片和caption
        all_paragraphs = list(self.doc.paragraphs)
        for paragraph in all_paragraphs:
            text = paragraph.text.strip()

            # 仅处理 Obsidian 图片语法，标准 Markdown 图片交由 pandoc 处理
            obsidian_match = obsidian_image_pattern.search(text)

            if obsidian_match:
                # 检查是否在同一段落包含caption
                caption_in_same_para = None
                if obsidian_match:
                    # 提取图片语法后的文本作为 potential caption
                    image_syntax_end = obsidian_match.end()
                    remaining_text = text[image_syntax_end:].strip()
                    image_path = obsidian_match.group(1)
                    # 检查剩余文本是否为caption
                    if remaining_text and caption_pattern.match(remaining_text):
                        caption_in_same_para = remaining_text

                # 查找图片实际路径
                actual_path = self._find_image_actual_path(image_path)
                if actual_path is None:
                    raise ImageProcessingError(f'Image referenced by Obsidian syntax was not found: {image_path}')

                # 构建图片信息
                image_info = {
                    'path': actual_path,
                    'paragraph': paragraph,
                    'caption_in_same_para': caption_in_same_para,
                }

                images_to_process.append(image_info)

        # 第二遍：处理图片插入
        for image_info in images_to_process:
            self._replace_paragraph_with_image(image_info['paragraph'], image_info)

        # 第三遍：处理caption格式化
        self._process_captions()

    def _find_image_actual_path(self, image_path: str) -> str | None:
        """查找图片的实际路径"""
        # Absolute local paths do not require search-path resolution.
        if os.path.isabs(image_path):
            return image_path if os.path.isfile(image_path) else None

        # Remote Obsidian embeds are intentionally unsupported because
        # python-docx cannot insert them without a separate download policy.
        if image_path.startswith(('http://', 'https://')):
            return None

        # 构建搜索路径列表
        search_paths = [
            self.source_dir,
            self.doc_dir,
            *DocumentConfig.get_image_search_paths(),  # 然后在配置的搜索路径中查找
        ]

        # 支持的图片格式
        supported_formats = DocumentConfig.IMAGE_CONFIG['supported_formats']

        for search_path_str in search_paths:
            search_path = Path(search_path_str).expanduser().resolve()
            if not search_path.exists() or not search_path.is_dir():
                continue

            try:
                # 构建目标图片路径
                image_path_obj = (search_path / image_path).resolve()

                # 直接匹配文件名
                if image_path_obj.is_file():
                    return str(image_path_obj)

                # 如果没有扩展名，尝试添加支持的格式
                if not image_path_obj.suffix:
                    for ext in supported_formats:
                        path_with_ext = image_path_obj.with_suffix(ext)
                        if path_with_ext.is_file():
                            return str(path_with_ext)

            except (OSError, ValueError):
                continue

        return None

    def _replace_paragraph_with_image(self, paragraph, image_info: dict):
        """将包含图片语法的段落替换为实际图片"""
        if not image_info['path'] or not os.path.exists(image_info['path']):
            raise ImageProcessingError(f'Image file is unavailable: {image_info["path"]}')

        if self._has_math_formula(paragraph):
            self._insert_image_before_math_caption(paragraph, image_info)
            return

        if image_info.get('caption_in_same_para'):
            p_element = paragraph._element
            parent = p_element.getparent()
            caption_p = self.doc.add_paragraph()
            caption_p.text = image_info['caption_in_same_para']
            parent.insert(parent.index(p_element) + 1, caption_p._element)

        paragraph.clear()
        run = paragraph.add_run()
        run.add_picture(image_info['path'], width=Inches(5))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        drawing_elements = run._element.xpath('.//w:drawing')
        if not drawing_elements:
            raise ImageProcessingError(f'Inserted image has no drawing element: {image_info["path"]}')
        self.image_formatter._set_image_wrap(drawing_elements[0])

    def _remove_image_syntax_only(self, paragraph):
        """只移除图片语法部分，保留其他内容（包括数学公式）"""
        original_text = paragraph.text
        cleaned_text = Patterns.OBSIDIAN_IMAGE_CLEANUP_PATTERN.sub('', original_text)
        cleaned_text = Patterns.MARKDOWN_IMAGE_CLEANUP_PATTERN.sub('', cleaned_text)
        cleaned_text = Patterns.WHITESPACE_CLEANUP_PATTERN.sub(' ', cleaned_text).strip()

        if cleaned_text == original_text:
            return

        for run in paragraph.runs:
            if run.text and ('![[' in run.text or '![' in run.text):
                new_run_text = Patterns.OBSIDIAN_IMAGE_CLEANUP_PATTERN.sub('', run.text)
                new_run_text = Patterns.MARKDOWN_IMAGE_CLEANUP_PATTERN.sub('', new_run_text)
                run.text = Patterns.WHITESPACE_CLEANUP_PATTERN.sub(' ', new_run_text).strip()

    def _insert_image_before_math_caption(self, caption_paragraph, image_info: dict):
        """在包含数学公式的caption前插入图片段落"""
        p_element = caption_paragraph._element
        parent = p_element.getparent()
        image_p = self.doc.add_paragraph()
        run = image_p.add_run()
        run.add_picture(image_info['path'], width=Inches(5))
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        drawing_elements = run._element.xpath('.//w:drawing')
        if not drawing_elements:
            raise ImageProcessingError(f'Inserted image has no drawing element: {image_info["path"]}')
        self.image_formatter._set_image_wrap(drawing_elements[0])
        parent.insert(parent.index(p_element), image_p._element)
        self._remove_image_syntax_only(caption_paragraph)

        cleaned_text = caption_paragraph.text.strip()
        if cleaned_text:
            self._processed_math_caption_texts.add(cleaned_text)
        self.image_formatter._format_image_caption(caption_paragraph)

    def _process_captions(self):
        """格式化所有图片和表格caption（位置已在预处理阶段调整）"""
        caption_pattern = Patterns.CAPTION_PATTERN
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text or text in self._processed_math_caption_texts:
                continue

            caption_match = caption_pattern.match(text)
            if not caption_match:
                continue

            caption_type = caption_match.group(1)
            number = caption_match.group(2)
            content = caption_match.group(3)
            has_math = self._has_math_formula(paragraph)

            if not has_math:
                prefix = '图' if caption_type in ['图', '图片', '图表'] else '表'
                paragraph.text = f'{prefix}{number}. {content}'

            if caption_type in ['图', '图片', '图表']:
                self.image_formatter._format_image_caption(paragraph)
            else:
                self.table_formatter.format_caption(paragraph)
