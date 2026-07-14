import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from src.config import DocumentConfig
from src.formatters import PageFormatter


class TestPageFormatter(unittest.TestCase):
    def test_setup_page_format_serializes_a4_portrait_for_every_section(self):
        document = Document()
        document.add_section(WD_SECTION.NEW_PAGE)

        for section in document.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)

        PageFormatter(DocumentConfig()).setup_page_format(document)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / 'a4.docx'
            document.save(output_path)
            saved_document = Document(output_path)

        self.assertEqual(2, len(saved_document.sections))
        for section in saved_document.sections:
            self.assertEqual(WD_ORIENT.PORTRAIT, section.orientation)
            self.assertAlmostEqual(210, section.page_width.mm, delta=0.01)
            self.assertAlmostEqual(297, section.page_height.mm, delta=0.01)
            self.assertAlmostEqual(34.58, section.top_margin.mm, delta=0.01)
            self.assertAlmostEqual(32.58, section.bottom_margin.mm, delta=0.01)
            self.assertAlmostEqual(28, section.left_margin.mm, delta=0.01)
            self.assertAlmostEqual(26, section.right_margin.mm, delta=0.01)

    def test_content_width_is_derived_from_page_size_and_margins(self):
        config = DocumentConfig()
        self.assertEqual(Mm(156), config.get_content_width_emu())

    def test_content_width_honors_instance_configuration(self):
        config = DocumentConfig()
        config.PAGE_SIZE = {'width': Mm(250), 'height': Mm(297)}
        config.PAGE_MARGINS = {'top': Mm(37), 'bottom': Mm(35), 'left': Mm(20), 'right': Mm(20)}

        self.assertEqual(Mm(210), config.get_content_width_emu())

    def test_document_grid_uses_configured_content_width(self):
        class CustomPageConfig(DocumentConfig):
            PAGE_SIZE = {'width': Mm(180), 'height': Mm(297)}
            PAGE_MARGINS = {'top': Mm(37), 'bottom': Mm(35), 'left': Mm(20), 'right': Mm(20)}
            CHARS_PER_LINE = 20

        config = CustomPageConfig()
        document = Document()
        PageFormatter(config).setup_page_format(document)

        document_grid = document.sections[0]._sectPr.find(qn('w:docGrid'))
        expected_char_space = round(Mm(140).twips / config.CHARS_PER_LINE)
        self.assertEqual(str(expected_char_space), document_grid.get(qn('w:charSpace')))

    def test_document_grid_uses_explicit_body_line_pitch(self):
        config = DocumentConfig()
        document = Document()

        PageFormatter(config).setup_page_format(document)

        document_grid = document.sections[0]._sectPr.find(qn('w:docGrid'))
        expected_line_pitch = config.BODY_LINE_PITCH.twips
        self.assertEqual(str(expected_line_pitch), document_grid.get(qn('w:linePitch')))

    def test_page_numbers_use_mirrored_odd_and_even_footers(self):
        config = DocumentConfig()
        document = Document()

        PageFormatter(config).add_page_numbers(document)

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / 'page-numbers.docx'
            document.save(output_path)
            document = Document(output_path)

        section = document.sections[0]
        odd_paragraph = section.footer.paragraphs[0]
        even_paragraph = section.even_page_footer.paragraphs[0]
        self.assertTrue(document.settings.odd_and_even_pages_header_footer)
        self.assertEqual(WD_PARAGRAPH_ALIGNMENT.RIGHT, odd_paragraph.alignment)
        self.assertEqual(WD_PARAGRAPH_ALIGNMENT.LEFT, even_paragraph.alignment)
        self.assertEqual(Pt(16), odd_paragraph.paragraph_format.right_indent)
        self.assertEqual(Pt(16), even_paragraph.paragraph_format.left_indent)
        odd_run_fonts = odd_paragraph.runs[0]._element.rPr.find(qn('w:rFonts'))
        self.assertEqual('Times New Roman', odd_run_fonts.get(qn('w:ascii')))
        self.assertEqual('Times New Roman', odd_run_fonts.get(qn('w:hAnsi')))
        self.assertEqual('SimSun', odd_run_fonts.get(qn('w:eastAsia')))
        self.assertEqual('— 1 —', odd_paragraph.text)


if __name__ == '__main__':
    unittest.main()
