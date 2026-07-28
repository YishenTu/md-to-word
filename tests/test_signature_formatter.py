import unittest
from decimal import Decimal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.config import DocumentConfig
from src.formatters.signature_formatter import SignatureFormatter
from src.utils.constants import ControlTokens


class TestSignatureFormatter(unittest.TestCase):
    def setUp(self):
        self.config = DocumentConfig()
        self.formatter = SignatureFormatter(self.config)

    def test_date_lengths_use_the_documented_character_grid_formula(self):
        expected_indents = {
            '2026\u5e747\u67081\u65e5': Decimal('12.75'),
            '2026\u5e7410\u67081\u65e5': Decimal('12.25'),
            '2026\u5e7410\u670814\u65e5': Decimal('11.75'),
        }

        for document_date, expected_indent in expected_indents.items():
            with self.subTest(document_date=document_date):
                self.assertEqual(
                    expected_indent,
                    self.formatter._calculate_left_indent_chars(document_date),
                )

    def test_adds_adjacent_centered_paragraphs_with_shared_character_indent(self):
        document = Document()
        document.add_paragraph('Body')
        signatory = 'Example Authority'
        document_date = '2026\u5e747\u67081\u65e5'

        self.formatter.add_signature(document, signatory, document_date)

        self.assertEqual(
            ['Body', '', '', signatory, document_date],
            [paragraph.text for paragraph in document.paragraphs],
        )
        for paragraph in document.paragraphs[1:3]:
            self.assertEqual(0, paragraph.paragraph_format.space_before.pt)
            self.assertEqual(0, paragraph.paragraph_format.space_after.pt)
            self.assertTrue(paragraph.paragraph_format.keep_together)
            self.assertTrue(paragraph.paragraph_format.keep_with_next)
            snap_to_grid = paragraph._element.pPr.find(qn('w:snapToGrid'))
            self.assertEqual('true', snap_to_grid.get(qn('w:val')))

        signatory_paragraph, date_paragraph = document.paragraphs[-2:]
        for paragraph in (signatory_paragraph, date_paragraph):
            self.assertEqual(WD_ALIGN_PARAGRAPH.CENTER, paragraph.alignment)
            indent = paragraph._element.pPr.find(qn('w:ind'))
            self.assertEqual('1275', indent.get(qn('w:leftChars')))
            self.assertIsNotNone(paragraph.paragraph_format.left_indent)

            run_fonts = paragraph.runs[0]._element.rPr.find(qn('w:rFonts'))
            self.assertEqual(
                self.config.FONTS['latin'],
                run_fonts.get(qn('w:ascii')),
            )
            self.assertEqual(
                self.config.FONTS['fangsong'],
                run_fonts.get(qn('w:eastAsia')),
            )

        self.assertTrue(signatory_paragraph.paragraph_format.keep_with_next)
        self.assertFalse(date_paragraph.paragraph_format.keep_with_next)
        auto_space = date_paragraph._element.pPr.find(qn('w:autoSpaceDN'))
        self.assertEqual('true', auto_space.get(qn('w:val')))

    def test_replaces_an_anchor_without_moving_following_content(self):
        document = Document()
        document.add_paragraph('Body')
        anchor = document.add_paragraph(ControlTokens.SIGNATURE)
        document.add_paragraph('Attachment declaration')

        self.formatter.replace_signature_anchor(
            anchor,
            'Example Authority',
            '2026\u5e747\u67081\u65e5',
        )

        self.assertEqual(
            ['Body', '', '', 'Example Authority', '2026\u5e747\u67081\u65e5', 'Attachment declaration'],
            [paragraph.text for paragraph in document.paragraphs],
        )
        self.assertNotIn(ControlTokens.SIGNATURE, [paragraph.text for paragraph in document.paragraphs])


if __name__ == '__main__':
    unittest.main()
