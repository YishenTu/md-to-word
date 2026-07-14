import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

from md_to_word import convert_document

PANDOC_AVAILABLE = shutil.which('pandoc') is not None
PROJECT_ROOT = Path(__file__).resolve().parents[1]


@unittest.skipUnless(PANDOC_AVAILABLE, 'Pandoc is required for example tests')
class TestBundledExamples(unittest.TestCase):
    def test_self_contained_example_converts_with_all_local_assets(self):
        input_path = PROJECT_ROOT / 'examples' / '公文格式示例.md'

        with tempfile.TemporaryDirectory() as directory:
            output_path = Path(directory) / 'example.docx'
            convert_document(input_path, output_path)
            document = Document(output_path)

        anchors = document.element.body.xpath('.//wp:anchor')
        page_breaks = document.element.body.xpath('.//w:br[@w:type="page"]')
        texts = [paragraph.text for paragraph in document.paragraphs]

        self.assertEqual(2, len(anchors))
        self.assertEqual(2, len(page_breaks))
        self.assertGreaterEqual(len(document.tables), 1)
        attachment_text = '附件：1. 示例实施方案\n　　　2. 项目数据明细表\n　　　3. 项目验收检查清单'
        self.assertEqual(1, texts.count(attachment_text))
        attachment_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == attachment_text)
        self.assertEqual(2, len(attachment_paragraph._element.xpath('.//w:br')))
        self.assertEqual(WD_ALIGN_PARAGRAPH.LEFT, attachment_paragraph.alignment)
        self.assertEqual(Pt(32), attachment_paragraph.paragraph_format.left_indent)
        self.assertEqual(Pt(0), attachment_paragraph.paragraph_format.first_line_indent)

        attachment_index = next(
            index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == attachment_text
        )
        attachment_spacing = document.paragraphs[attachment_index - 1]
        self.assertEqual('', attachment_spacing.text)
        self.assertEqual(Pt(0), attachment_spacing.paragraph_format.space_before)
        self.assertEqual(Pt(0), attachment_spacing.paragraph_format.space_after)
        snap_to_grid = attachment_spacing._element.pPr.find(qn('w:snapToGrid'))
        self.assertIsNotNone(snap_to_grid)
        self.assertEqual('true', snap_to_grid.get(qn('w:val')))
        self.assertNotEqual('', document.paragraphs[attachment_index - 2].text)
        signatory_text = '\u793a\u4f8b\u67d0\u67d0\u5730\u533a\u67d0\u67d0\u5b57\u53f7\u6709\u9650\u516c\u53f8'
        document_date = '2026\u5e747\u670814\u65e5'
        signature_spacing = document.paragraphs[attachment_index + 1 : attachment_index + 3]
        self.assertEqual(['', ''], [paragraph.text for paragraph in signature_spacing])
        for paragraph in signature_spacing:
            snap_to_grid = paragraph._element.pPr.find(qn('w:snapToGrid'))
            self.assertEqual('true', snap_to_grid.get(qn('w:val')))

        self.assertEqual(signatory_text, document.paragraphs[attachment_index + 3].text)
        self.assertEqual(document_date, document.paragraphs[attachment_index + 4].text)
        for paragraph in document.paragraphs[attachment_index + 3 : attachment_index + 5]:
            self.assertEqual(WD_ALIGN_PARAGRAPH.CENTER, paragraph.alignment)
            indent = paragraph._element.pPr.find(qn('w:ind'))
            self.assertEqual('1225', indent.get(qn('w:leftChars')))
        self.assertEqual(
            'true',
            document.paragraphs[attachment_index + 4]._element.pPr.find(qn('w:autoSpaceDN')).get(qn('w:val')),
        )
        self.assertIn('图1. 项目实施流程 Project workflow', texts)
        self.assertIn('图2. 项目评估与数据复核 Project review', texts)

        image_paragraph_indices = [
            index for index, paragraph in enumerate(document.paragraphs) if paragraph._element.xpath('.//wp:anchor')
        ]
        self.assertEqual(
            {
                '图1. 项目实施流程 Project workflow',
                '图2. 项目评估与数据复核 Project review',
            },
            {document.paragraphs[index + 1].text for index in image_paragraph_indices},
        )
        for index in image_paragraph_indices:
            paragraph_format = document.paragraphs[index].paragraph_format
            self.assertEqual(Pt(1), paragraph_format.line_spacing)
            self.assertEqual(WD_LINE_SPACING.EXACTLY, paragraph_format.line_spacing_rule)

        for caption_text in (
            '图1. 项目实施流程 Project workflow',
            '图2. 项目评估与数据复核 Project review',
        ):
            caption = next(paragraph for paragraph in document.paragraphs if paragraph.text == caption_text)
            run_fonts = caption.runs[0]._element.rPr.find(qn('w:rFonts'))
            self.assertEqual('Times New Roman', run_fonts.get(qn('w:ascii')))
            self.assertEqual('FangSong', run_fonts.get(qn('w:eastAsia')))


if __name__ == '__main__':
    unittest.main()
