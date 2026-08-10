import shutil
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from md_to_word import convert_document, resolve_output_path
from src.utils.exceptions import ImageProcessingError


class TestCliConversion(unittest.TestCase):
    def test_legacy_doc_extension_is_normalized_to_docx(self):
        input_path = Path('/tmp/report.md')

        output_path = resolve_output_path(input_path, '/tmp/report.doc')

        self.assertEqual(Path('/tmp').resolve() / 'report.docx', output_path)

    def test_failed_postprocessing_preserves_existing_output(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            input_path = root / 'input.md'
            output_path = root / 'output.docx'
            input_path.write_text('正文', encoding='utf-8')
            original_output = b'previous valid document'
            output_path.write_bytes(original_output)

            def create_raw_docx(_content, staging_path, **_kwargs):
                Document().save(staging_path)
                return staging_path

            with patch(
                'md_to_word.PandocProcessor.check_pandoc_available',
                return_value=True,
            ):
                with patch(
                    'md_to_word.PandocProcessor.convert_markdown_to_docx',
                    side_effect=create_raw_docx,
                ):
                    with patch(
                        'md_to_word.WordPostprocessor.apply_formatting',
                        side_effect=ImageProcessingError('missing image'),
                    ):
                        with self.assertRaises(ImageProcessingError):
                            convert_document(input_path, output_path)

            self.assertEqual(original_output, output_path.read_bytes())
            self.assertEqual(
                {'input.md', 'output.docx'},
                {path.name for path in root.iterdir()},
            )

    @unittest.skipUnless(shutil.which('pandoc'), 'Pandoc is required for CLI integration tests')
    def test_successful_conversion_atomically_publishes_a_valid_docx(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            input_path = root / 'input.md'
            output_path = root / 'output.docx'
            input_path.write_text('正文。', encoding='utf-8')

            result = convert_document(input_path, output_path)

            document = Document(result)
            self.assertEqual('input', document.paragraphs[0].text)
            self.assertEqual(
                {'input.md', 'output.docx'},
                {path.name for path in root.iterdir()},
            )

    @unittest.skipUnless(shutil.which('pandoc'), 'Pandoc is required for CLI integration tests')
    def test_single_h1_overrides_filename_as_the_formatted_document_title(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            input_path = root / 'filename-title.md'
            output_path = root / 'output.docx'
            input_path.write_text('# Markdown Title\n\nBody.', encoding='utf-8')

            convert_document(input_path, output_path)
            document = Document(output_path)

        title = document.paragraphs[0]
        self.assertEqual('Markdown Title', title.text)
        self.assertNotIn('# Markdown Title', [paragraph.text for paragraph in document.paragraphs])
        self.assertEqual(Pt(22), title.runs[0].font.size)
        run_fonts = title.runs[0]._element.rPr.find(qn('w:rFonts'))
        self.assertEqual('FZXiaoBiaoSong-B05S', run_fonts.get(qn('w:eastAsia')))


if __name__ == '__main__':
    unittest.main()
