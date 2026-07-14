import base64
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt

from src.config import DocumentConfig
from src.core.markdown_preprocessor import MarkdownPreprocessor
from src.core.pandoc_processor import PandocProcessor
from src.core.word_postprocessor import WordPostprocessor
from src.utils.exceptions import ImageProcessingError, PandocError

PANDOC_AVAILABLE = shutil.which('pandoc') is not None
ONE_PIXEL_PNG = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII='
)


@unittest.skipUnless(PANDOC_AVAILABLE, 'Pandoc is required for pipeline tests')
class TestConversionPipeline(unittest.TestCase):
    def _convert(self, markdown, directory, source_dir=None, title='审计测试'):
        preprocessor = MarkdownPreprocessor()
        content = preprocessor.preprocess_content(markdown)
        metadata = {
            'title': title,
        }
        output = Path(directory) / 'output.docx'
        PandocProcessor().convert_markdown_to_docx(content, str(output))
        WordPostprocessor().apply_formatting(
            str(output),
            metadata,
            source_dir=source_dir or directory,
        )
        return Document(output)

    def assert_run_fonts(self, run, east_asia_font):
        run_fonts = run._element.rPr.find(qn('w:rFonts'))
        self.assertIsNotNone(run_fonts)
        self.assertEqual('Times New Roman', run_fonts.get(qn('w:ascii')))
        self.assertEqual('Times New Roman', run_fonts.get(qn('w:hAnsi')))
        self.assertEqual(east_asia_font, run_fonts.get(qn('w:eastAsia')))

    def assert_paragraph_fonts(self, paragraph, east_asia_font):
        text_runs = [run for run in paragraph.runs if run.text]
        self.assertTrue(text_runs)
        for run in text_runs:
            self.assert_run_fonts(run, east_asia_font)

    def test_latin_text_uses_times_new_roman_at_every_content_level(self):
        markdown = (
            '## Heading 一级\n\n'
            '### Subheading 二级\n\n'
            'Body 正文 2026\n\n'
            '- Item 列表\n\n'
            '| Column 列 |\n| --- |\n| Value 值 |'
        )

        with tempfile.TemporaryDirectory() as directory:
            document = self._convert(
                markdown,
                directory,
                title='Report 公文 2026',
            )

        paragraphs_by_text = {paragraph.text: paragraph for paragraph in document.paragraphs if paragraph.text}
        fonts = DocumentConfig.FONTS
        self.assert_paragraph_fonts(
            paragraphs_by_text['Report 公文 2026'],
            fonts['xiaobiaosong'],
        )
        self.assert_paragraph_fonts(paragraphs_by_text['Heading 一级'], fonts['heiti'])
        self.assert_paragraph_fonts(paragraphs_by_text['Subheading 二级'], fonts['kaiti'])
        self.assert_paragraph_fonts(paragraphs_by_text['Body 正文 2026'], fonts['fangsong'])
        self.assert_paragraph_fonts(paragraphs_by_text['Item 列表'], fonts['fangsong'])
        self.assert_paragraph_fonts(
            document.tables[0].cell(1, 0).paragraphs[0],
            fonts['fangsong'],
        )
        self.assert_paragraph_fonts(
            document.sections[0].footer.paragraphs[0],
            fonts['songti'],
        )

        numbering_fonts = document.part.numbering_part._element.xpath('.//w:lvl/w:rPr/w:rFonts')
        self.assertTrue(numbering_fonts)
        for run_fonts in numbering_fonts:
            self.assertEqual('Times New Roman', run_fonts.get(qn('w:ascii')))
            self.assertEqual('Times New Roman', run_fonts.get(qn('w:hAnsi')))
            self.assertEqual(fonts['fangsong'], run_fonts.get(qn('w:eastAsia')))

    def test_unordered_lists_use_controlled_native_bullet_layout(self):
        markdown = (
            '- Top-level item with enough content to exercise wrapped-line alignment in rendered output.\n'
            '  - Nested item\n'
            '- Final item'
        )

        with tempfile.TemporaryDirectory() as directory:
            document = self._convert(markdown, directory)

        paragraphs = {
            paragraph.text: paragraph
            for paragraph in document.paragraphs
            if paragraph.text
            in {
                'Top-level item with enough content to exercise wrapped-line alignment in rendered output.',
                'Nested item',
                'Final item',
            }
        }
        self.assertEqual(3, len(paragraphs))

        expected_layout = {
            'Top-level item with enough content to exercise wrapped-line alignment in rendered output.': (
                '0',
                '960',
            ),
            'Nested item': ('1', '1600'),
            'Final item': ('0', '960'),
        }
        for text, (expected_level, expected_left) in expected_layout.items():
            paragraph = paragraphs[text]
            p_pr = paragraph._element.find(qn('w:pPr'))
            num_pr = p_pr.find(qn('w:numPr'))
            self.assertIsNotNone(num_pr)
            self.assertEqual(expected_level, num_pr.find(qn('w:ilvl')).get(qn('w:val')))

            indent = p_pr.find(qn('w:ind'))
            self.assertEqual(expected_left, indent.get(qn('w:left')))
            self.assertEqual('320', indent.get(qn('w:hanging')))
            tab = p_pr.find(qn('w:tabs')).find(qn('w:tab'))
            self.assertEqual('num', tab.get(qn('w:val')))
            self.assertEqual(expected_left, tab.get(qn('w:pos')))
            self.assertEqual(Pt(0), paragraph.paragraph_format.space_before)
            self.assertEqual(Pt(0), paragraph.paragraph_format.space_after)
            self.assertNotIn('•', paragraph.text)

        bullet_levels = document.part.numbering_part._element.xpath('.//w:lvl[w:numFmt[@w:val="bullet"]]')
        self.assertTrue(bullet_levels)
        for level in bullet_levels:
            self.assertEqual('•', level.find(qn('w:lvlText')).get(qn('w:val')))
            self.assertEqual('tab', level.find(qn('w:suff')).get(qn('w:val')))
            self.assertEqual('32', level.find(qn('w:rPr')).find(qn('w:sz')).get(qn('w:val')))

    def test_standard_markdown_images_use_the_shared_image_formatter(self):
        encoded_image = base64.b64encode(ONE_PIXEL_PNG).decode('ascii')
        markdown = f'![Architecture diagram](data:image/png;base64,{encoded_image})\n\n图1: Architecture'

        with tempfile.TemporaryDirectory() as directory:
            document = self._convert(markdown, directory)

        anchors = document.element.body.xpath('.//wp:anchor')
        inlines = document.element.body.xpath('.//wp:inline')
        extents = document.element.body.xpath('.//wp:anchor/wp:extent')
        self.assertEqual(1, len(anchors))
        self.assertEqual(0, len(inlines))
        self.assertEqual(str(int(DocumentConfig().get_content_width_emu())), extents[0].get('cx'))
        image_properties = anchors[0].xpath('./wp:docPr')[0]
        self.assertEqual('Architecture diagram', image_properties.get('descr'))
        self.assertNotIn('Architecture diagram', [paragraph.text for paragraph in document.paragraphs])
        self.assertEqual('', document.paragraphs[1].text)

        image_paragraph_index = next(
            index for index, paragraph in enumerate(document.paragraphs) if paragraph._element.xpath('.//wp:anchor')
        )
        image_paragraph = document.paragraphs[image_paragraph_index]
        caption_paragraph = document.paragraphs[image_paragraph_index + 1]
        self.assertEqual('图1. Architecture', caption_paragraph.text)
        self.assert_paragraph_fonts(caption_paragraph, 'FangSong')
        self.assertEqual(Pt(1), image_paragraph.paragraph_format.line_spacing)
        self.assertEqual(
            WD_LINE_SPACING.EXACTLY,
            image_paragraph.paragraph_format.line_spacing_rule,
        )
        self.assertEqual(Pt(0), caption_paragraph.paragraph_format.space_before)

    def test_attachment_declaration_is_not_duplicated(self):
        markdown = '正文。\n\n附件：1. 实施方案'

        with tempfile.TemporaryDirectory() as directory:
            document = self._convert(markdown, directory)

        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertEqual(1, texts.count('附件：1. 实施方案'))

    def test_soft_breaks_and_fenced_code_preserve_text_content(self):
        markdown = 'The first line\nends here\n\n```markdown\n**bold marker**\n* list-like code\n---\n```'

        with tempfile.TemporaryDirectory() as directory:
            document = self._convert(markdown, directory)

        text = '\n'.join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn('The first line ends here', text)
        self.assertIn('**bold marker**', text)
        self.assertIn('* list-like code', text)
        self.assertIn('---', text)

    def test_table_text_is_normalized_without_preprocessing_source_markup(self):
        markdown = '| Column |\n| --- |\n| **Value** |'

        with tempfile.TemporaryDirectory() as directory:
            document = self._convert(markdown, directory)

        value_paragraph = document.tables[0].cell(1, 0).paragraphs[0]
        self.assertEqual('Value', value_paragraph.text)
        self.assertTrue(all(run.bold is False for run in value_paragraph.runs))

    def test_obsidian_images_resolve_relative_to_the_markdown_source(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source_dir = root / 'source'
            output_dir = root / 'output'
            source_dir.mkdir()
            output_dir.mkdir()
            (source_dir / 'diagram.png').write_bytes(ONE_PIXEL_PNG)

            document = self._convert('![[diagram.png]]', output_dir, source_dir)

        self.assertEqual(1, len(document.element.body.xpath('.//wp:anchor')))

    def test_missing_obsidian_images_fail_the_conversion(self):
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaises(ImageProcessingError):
                self._convert('![[missing.png]]', directory)

    def test_missing_standard_markdown_images_fail_the_conversion(self):
        with tempfile.TemporaryDirectory() as directory:
            with self.assertRaises(PandocError):
                self._convert('![](missing.png)', directory)


if __name__ == '__main__':
    unittest.main()
