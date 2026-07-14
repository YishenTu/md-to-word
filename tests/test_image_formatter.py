import unittest
from unittest.mock import Mock

from docx import Document
from docx.oxml.shared import OxmlElement
from docx.shared import Mm

from src.config import DocumentConfig
from src.formatters import ImageFormatter
from src.utils.exceptions import ImageProcessingError


class TestImageFormatter(unittest.TestCase):
    def setUp(self):
        self.config = DocumentConfig()
        self.config.PAGE_SIZE = {'width': Mm(250), 'height': Mm(297)}
        self.config.PAGE_MARGINS = {'top': Mm(37), 'bottom': Mm(35), 'left': Mm(20), 'right': Mm(20)}
        self.formatter = ImageFormatter(self.config)

    def test_full_width_image_uses_configured_content_width(self):
        drawing = OxmlElement('w:drawing')
        inline = OxmlElement('wp:inline')
        inline_extent = OxmlElement('wp:extent')
        inline_extent.set('cx', '100')
        inline_extent.set('cy', '50')
        inline.append(inline_extent)
        drawing.append(inline)

        graphic = OxmlElement('a:graphic')
        graphic_data = OxmlElement('a:graphicData')
        picture = OxmlElement('pic:pic')
        shape_properties = OxmlElement('pic:spPr')
        transform = OxmlElement('a:xfrm')
        picture_extent = OxmlElement('a:ext')
        picture_extent.set('cx', '100')
        picture_extent.set('cy', '50')
        transform.append(picture_extent)
        shape_properties.append(transform)
        picture.append(shape_properties)
        graphic_data.append(picture)
        graphic.append(graphic_data)
        inline.append(graphic)

        self.formatter._set_image_full_width(drawing)

        expected_width = str(int(Mm(210)))
        expected_height = str(int(Mm(105)))
        self.assertEqual(expected_width, inline_extent.get('cx'))
        self.assertEqual(expected_height, inline_extent.get('cy'))
        self.assertEqual(expected_width, picture_extent.get('cx'))
        self.assertEqual(expected_height, picture_extent.get('cy'))

    def test_wrapped_image_uses_configured_content_width(self):
        parent = Mock()
        inline = Mock()
        inline.getparent.return_value = parent

        extent = Mock()
        extent.get.side_effect = lambda name: {'cx': '100', 'cy': '50'}[name]

        doc_property = OxmlElement('wp:docPr')
        doc_property.set('id', '7')
        doc_property.set('name', 'Diagram')
        doc_property.set('descr', 'Architecture diagram')

        graphic = Mock()
        drawing = Mock()
        self.formatter.xml_processor.process_image_properties = Mock(
            return_value={'inline': [inline], 'extent': [extent], 'docPr': [doc_property], 'graphic': [graphic]}
        )
        anchor = Mock()
        self.formatter._create_anchor_element = Mock(return_value=anchor)

        self.formatter._set_image_wrap(drawing)

        expected_width = str(int(Mm(210)))
        expected_height = str(int(Mm(105)))
        self.formatter._create_anchor_element.assert_called_once_with(
            expected_width,
            expected_height,
            doc_property,
        )
        anchor.append.assert_called_once_with(graphic)
        parent.replace.assert_called_once_with(inline, anchor)

    def test_image_xml_failures_are_not_silently_ignored(self):
        document = Mock()
        document.paragraphs = [Mock()]
        self.formatter.xml_processor.find_drawings_in_paragraphs = Mock(return_value={0: [Mock()]})
        self.formatter.xml_processor.process_image_properties = Mock(side_effect=ValueError('invalid drawing'))

        with self.assertRaises(ImageProcessingError):
            self.formatter.format_images(document)

    def test_image_paragraph_with_visible_text_is_not_collapsed(self):
        document = Document()
        paragraph = document.add_paragraph('Visible text')
        drawing = OxmlElement('w:drawing')
        drawing.append(OxmlElement('wp:anchor'))
        paragraph._element.append(drawing)

        self.formatter._collapse_anchored_image_paragraph(paragraph)

        self.assertIsNone(paragraph.paragraph_format.line_spacing)


if __name__ == '__main__':
    unittest.main()
